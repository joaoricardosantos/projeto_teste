"""
execucao_api.py
Endpoints para protocolar ações de execução de taxas condominiais.
Fluxo:
  1. Frontend busca inadimplentes de um condomínio (já existe em condominio_service)
  2. Usuário seleciona unidades e preenche dados do condomínio / partes
  3. Backend gera DOCX e PDF para cada unidade selecionada (zip com todos)
"""
import io
import os
import re
import zipfile
import logging
from datetime import date

_MODELO_DOCX             = os.path.join(os.path.dirname(__file__), "..", "..", "frontend", "public", "modelo_execucao.docx")
_MODELO_DOCX_HONORARIOS  = os.path.join(os.path.dirname(__file__), "..", "..", "frontend", "public", "modelo_execucao_honorarios.docx")

def _modelo_path(modelo: str) -> str:
    return _MODELO_DOCX_HONORARIOS if modelo == "com_honorarios" else _MODELO_DOCX

from django.http import HttpResponse
from ninja import Router
from ninja.errors import HttpError

from core.auth import JWTAuth
from core.condominio_service import get_unidades_inadimplentes
from core.superlogica import verificar_condominio, _get_sl, _get_headers
from django.conf import settings

logger = logging.getLogger(__name__)
execucao_router = Router(auth=JWTAuth())


# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

_ESTADOS_SL = {
    "1":"AC","2":"AL","3":"AM","4":"AP","5":"BA","6":"CE","7":"DF","8":"ES",
    "9":"GO","10":"MA","11":"MG","12":"MS","13":"MT","14":"PA","15":"PB",
    "16":"PE","17":"PI","18":"PR","19":"RJ","20":"RN","21":"RO","22":"RR",
    "23":"RS","24":"SC","25":"SE","26":"SP","27":"TO",
}


def _buscar_responsaveis_por_unidade(id_condominio: int) -> dict:
    """Retorna mapa {id_unidade: {email, endereco}} buscando todos os responsáveis do condomínio."""
    mapa = {}
    pagina = 1
    while True:
        try:
            resp = _get_sl(
                f"{settings.SUPERLOGICA_BASE_URL}/responsaveis",
                params={"idCondominio": id_condominio, "pagina": pagina, "itensPorPagina": 50},
                timeout=30,
            )
            if not resp.ok:
                break
            dados = resp.json()
            if not dados:
                break
            if isinstance(dados, dict):
                dados = [dados]
            for item in dados:
                id_uni = item.get("id_unidade_uni")
                if not id_uni:
                    continue
                # só preenche se ainda não tiver (prioriza primeiro responsável)
                if id_uni in mapa:
                    continue
                email = (item.get("st_email_con") or "").strip()
                # Monta endereço completo
                partes = [
                    (item.get("st_endereco_con") or "").strip(),
                    (item.get("st_complemento_con") or "").strip(),
                    (item.get("st_bairro_con") or "").strip(),
                ]
                endereco = ", ".join(p for p in partes if p)
                cidade = (item.get("st_cidade_con") or "").strip()
                uf     = _ESTADOS_SL.get(str(item.get("st_estado_con", "")), "")
                cep    = (item.get("st_cep_con") or "").strip()
                if cidade and uf:
                    endereco += f", {cidade}/{uf}"
                elif cidade:
                    endereco += f", {cidade}"
                if cep:
                    endereco += f" - CEP {cep[:5]}-{cep[5:]}" if len(cep) == 8 else f" - CEP {cep}"
                mapa[id_uni] = {
                    "email":    email.split(";")[0].strip() if email else "",
                    "endereco": endereco.strip(", "),
                }
            if len(dados) < 50:
                break
            pagina += 1
        except Exception as e:
            logger.warning(f"Erro ao buscar responsáveis do condomínio {id_condominio}: {e}")
            break
    return mapa


def _buscar_sindico(id_condominio: int) -> dict:
    """Busca o síndico ativo do condomínio via /sindicos."""
    try:
        resp = _get_sl(
            f"{settings.SUPERLOGICA_BASE_URL}/sindicos",
            params={"idCondominio": id_condominio, "pagina": 1, "itensPorPagina": 50},
            timeout=20,
        )
        if not resp.ok:
            return {}
        dados = resp.json()
        if not dados:
            return {}
        if isinstance(dados, dict):
            dados = [dados]
        # Prioriza quem tem cargo "Síndico" e ainda está na gestão (dt_saida vazia ou futura)
        from datetime import datetime
        hoje = datetime.today()
        def _esta_ativo(s):
            saida = s.get("dt_saida_sin", "")
            if not saida:
                return True
            try:
                return datetime.strptime(saida[:10], "%m/%d/%Y") >= hoje
            except Exception:
                return True

        sindicos = [s for s in dados if _esta_ativo(s)]
        # Preferência: cargo "Síndico" > qualquer outro
        principal = next(
            (s for s in sindicos if "ndico" in (s.get("st_cargo_sin") or "")),
            sindicos[0] if sindicos else None,
        )
        if not principal:
            return {}
        cpf_raw = (principal.get("st_cpf_sin") or "").strip().replace(".", "").replace("-", "")
        # Formata CPF: 00000000000 -> 000.000.000-00
        if len(cpf_raw) == 11 and cpf_raw.isdigit():
            cpf_fmt = f"{cpf_raw[:3]}.{cpf_raw[3:6]}.{cpf_raw[6:9]}-{cpf_raw[9:]}"
        else:
            cpf_fmt = cpf_raw
        return {
            "nome_sindico": (principal.get("st_nome_sin") or "").strip(),
            "cpf_sindico":  cpf_fmt,
            "cargo_sindico": (principal.get("st_cargo_sin") or "").strip(),
        }
    except Exception as e:
        logger.warning(f"Erro ao buscar síndico do condomínio {id_condominio}: {e}")
        return {}


def _buscar_dados_condominio(id_condominio: int) -> dict:
    """Busca CNPJ, endereço, síndico via /configuracoes + /sindicos em paralelo."""
    import concurrent.futures

    def _cfg():
        try:
            resp = _get_sl(
                f"{settings.SUPERLOGICA_BASE_URL}/configuracoes",
                params={"idCondominio": id_condominio, "pagina": 1, "itensPorPagina": 1},
                timeout=20,
            )
            if not resp.ok:
                return {}
            dados = resp.json()
            if not dados:
                return {}
            d = dados[0] if isinstance(dados, list) else dados
            endereco_parts = [
                d.get("st_endereco_cond", ""),
                d.get("st_numeroendereco_cond", ""),
                d.get("st_complemento_cond", ""),
                d.get("st_bairro_cond", ""),
            ]
            endereco = ", ".join(p for p in endereco_parts if p and p.strip())
            cidade = d.get("st_cidade_cond", "")
            uf = _ESTADOS_SL.get(str(d.get("st_estado_cond", "")), d.get("st_estado_cond", ""))
            if cidade and uf:
                endereco += f", {cidade}/{uf}"
            elif cidade:
                endereco += f", {cidade}"
            # st_cnpj_cond pode estar vazio; Superlógica às vezes usa st_cpf_cond para CNPJ de PJ
            cnpj_bruto = (d.get("st_cnpj_cond") or d.get("st_cpf_cond") or "").strip()
            cnpj_raw = cnpj_bruto.replace(".", "").replace("/", "").replace("-", "")
            if len(cnpj_raw) == 14 and cnpj_raw.isdigit():
                cnpj_fmt = f"{cnpj_raw[:2]}.{cnpj_raw[2:5]}.{cnpj_raw[5:8]}/{cnpj_raw[8:12]}-{cnpj_raw[12:]}"
            else:
                cnpj_fmt = cnpj_bruto
            return {
                "nome_condominio": (d.get("st_nome_cond") or "").strip(),
                "cnpj":            cnpj_fmt,
                "endereco":        endereco.strip(", "),
            }
        except Exception as e:
            logger.warning(f"Erro ao buscar configurações do condomínio {id_condominio}: {e}")
            return {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as ex:
        f_cfg     = ex.submit(_cfg)
        f_sindico = ex.submit(_buscar_sindico, id_condominio)
        cfg     = f_cfg.result()
        sindico = f_sindico.result()

    return {**cfg, **sindico}


def _fmt_brl(valor) -> str:
    try:
        v = float(str(valor).replace("R$", "").replace(".", "").replace(",", ".").strip())
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(valor)


def _substituir(texto: str, variaveis: dict) -> str:
    """Substitui todos os [CAMPOS] pelo valor correspondente."""
    for chave, valor in variaveis.items():
        texto = texto.replace(f"[{chave}]", valor or "")
    return texto


def _valor_para_modelo(unidade: dict, modelo: str) -> str:
    """
    Retorna o valor formatado conforme o modelo:
    - 'sem_honorarios': valor total menos honorários
    - qualquer outro ('com_honorarios'): valor total
    """
    if modelo == "sem_honorarios":
        return unidade.get("valor_sem_honorarios") or unidade.get("valor", "")
    return unidade.get("valor", "")


def _variaveis(dados_condo: dict, dados_unidade: dict, responsavel: dict | None = None, modelo: str = "com_honorarios") -> dict:
    return {
        # Condomínio
        "NOME DO CONDOMÍNIO": dados_condo.get("nome_condominio", ""),
        "CNPJ":               dados_condo.get("cnpj", ""),
        "endereço":           dados_condo.get("endereco", ""),
        "nome do síndico":    dados_condo.get("nome_sindico", ""),
        "CPF do síndico":     dados_condo.get("cpf_sindico", ""),
        "NOME DA COMARCA":    dados_condo.get("comarca", "Natal"),
        "NOME DA CIDADE":     dados_condo.get("comarca", "Natal"),
        # Executado
        "NOME DA PARTE EXECUTADA":  dados_unidade.get("nome", ""),
        "CPF da parte executada":   dados_unidade.get("cpf", ""),
        "endereço":                 dados_unidade.get("endereco") or dados_condo.get("endereco", ""),
        "telefone":                 dados_unidade.get("telefone", ""),
        "endereço de e-mail":       dados_unidade.get("email", ""),
        # Débito
        "nome do condomínio":       dados_condo.get("nome_condominio", ""),
        "nº do apartamento":        dados_unidade.get("unidade", ""),
        "valor do débito":          _valor_para_modelo(dados_unidade, modelo),
        # Responsável pela petição
        "NOME DO RESPONSÁVEL PELA PETIÇÃO": (responsavel or {}).get("nome", ""),
        "Função": (responsavel or {}).get("funcao", ""),
    }


# ──────────────────────────────────────────────
# Geração DOCX
# ──────────────────────────────────────────────

def _gerar_docx(variaveis: dict, imagem_bytes: bytes | None = None, modelo_path: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    import copy

    doc = Document(modelo_path or _MODELO_DOCX)

    def _subst_para(para):
        if not para.runs:
            return

        # Passo 1: substitui campo a campo dentro de cada run individualmente.
        # Campos que estão inteiramente dentro de um run são resolvidos aqui,
        # preservando bold/italic/underline do run original.
        for run in para.runs:
            if run.text:
                run.text = _substituir(run.text, variaveis)

        # Passo 2: verifica se ainda há campos que cruzam runs
        texto_atual = "".join(r.text for r in para.runs)
        texto_final = _substituir(texto_atual, variaveis)
        if texto_final == texto_atual:
            return  # tudo resolvido no passo 1

        # Passo 3: colapsa apenas os grupos de runs que formam um campo cruzado.
        # Para cada campo ainda presente no texto, localiza os runs que o compõem
        # e colapsa apenas esse grupo, usando a formatação do run onde o '[' aparece.
        runs = list(para.runs)
        i = 0
        while i < len(runs):
            # Acumula runs até fechar um '[...]'
            if "[" not in runs[i].text:
                i += 1
                continue
            # Verifica se o campo está completo neste run
            if "]" in runs[i].text.split("[", 1)[1]:
                i += 1
                continue
            # Campo abre aqui mas fecha em outro run — acumula
            grupo_inicio = i
            fmt_bold      = runs[i].bold
            fmt_italic    = runs[i].italic
            fmt_underline = runs[i].underline
            fmt_size      = runs[i].font.size
            fmt_name      = runs[i].font.name
            texto_grupo   = runs[i].text
            j = i + 1
            while j < len(runs) and "]" not in texto_grupo:
                texto_grupo += runs[j].text
                j += 1
            # texto_grupo agora contém o campo completo (talvez com mais texto)
            texto_grupo_sub = _substituir(texto_grupo, variaveis)
            # Reescreve o primeiro run do grupo com o texto substituído
            runs[grupo_inicio].text      = texto_grupo_sub
            runs[grupo_inicio].bold      = fmt_bold
            runs[grupo_inicio].italic    = fmt_italic
            runs[grupo_inicio].underline = fmt_underline
            if fmt_size:
                runs[grupo_inicio].font.size = fmt_size
            if fmt_name:
                runs[grupo_inicio].font.name = fmt_name
            # Zera os runs consumidos pelo grupo
            for k in range(grupo_inicio + 1, j):
                runs[k].text = ""
            i = j

    def _processar_paragrafos(paragrafos):
        for para in paragrafos:
            _subst_para(para)

    _processar_paragrafos(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _processar_paragrafos(cell.paragraphs)

    # Substitui parágrafo [printscreen do Doc. 06] por imagem (se fornecida)
    if imagem_bytes:
        _inserir_imagem_placeholder(doc, imagem_bytes)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _inserir_imagem_placeholder(doc, imagem_bytes: bytes):
    """
    Encontra o parágrafo com texto '[printscreen do Doc. 06]' e substitui
    todos os seus runs por uma imagem inline centralizada.
    """
    from docx.shared import Inches
    from docx.oxml.ns import qn
    import lxml.etree as etree

    PLACEHOLDER = "[printscreen do Doc. 06]"

    for para in doc.paragraphs:
        # Texto completo do parágrafo (ignorando fragmentação de runs)
        texto_full = "".join(r.text for r in para.runs)
        if PLACEHOLDER not in texto_full:
            continue

        # Limpa todos os runs existentes
        for run in para.runs:
            run.text = ""

        # Usa o primeiro run para inserir a imagem
        run = para.runs[0] if para.runs else para.add_run()

        # Calcula largura: ocupa a área de texto da página (aprox. 15 cm = 5.9")
        img_buf = io.BytesIO(imagem_bytes)
        try:
            run.add_picture(img_buf, width=Inches(5.9))
        except Exception as e:
            # Fallback: insere como texto de erro
            run.text = f"[Imagem não pôde ser inserida: {e}]"
        break


# ──────────────────────────────────────────────
# Geração PDF via LibreOffice (converte DOCX preenchido)
# ──────────────────────────────────────────────

def _gerar_pdf(variaveis: dict, nome_condo: str = "", nome_unidade: str = "", imagem_bytes: bytes | None = None, modelo_path: str | None = None) -> bytes:
    import subprocess
    import tempfile

    docx_bytes = _gerar_docx(variaveis, imagem_bytes=imagem_bytes, modelo_path=modelo_path)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "documento.docx")
        pdf_path  = os.path.join(tmpdir, "documento.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, docx_path,
            ],
            capture_output=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice falhou: {result.stderr.decode()[:300]}")

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF não foi gerado pelo LibreOffice")

        with open(pdf_path, "rb") as f:
            return f.read()


# ──────────────────────────────────────────────
# Endpoints
# ──────────────────────────────────────────────

@execucao_router.get("/inadimplentes", response={200: dict})
def listar_inadimplentes(request, id_condominio: int):
    """Retorna inadimplentes + dados do condomínio para preenchimento automático."""
    from datetime import datetime
    from core.superlogica import buscar_unidades, buscar_inadimplentes_condominio

    try:
        # Busca tudo em paralelo
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as ex:
            f_condo  = ex.submit(_buscar_dados_condominio, id_condominio)
            f_inadim = ex.submit(get_unidades_inadimplentes, id_condominio)
            f_resp   = ex.submit(_buscar_responsaveis_por_unidade, id_condominio)
            f_mapa   = ex.submit(buscar_unidades, id_condominio)
            dados_condo  = f_condo.result()
            unidades     = f_inadim.result()
            resp_mapa    = f_resp.result()
            mapa         = f_mapa.result() or {}

        mapa.pop("__nome_cond__", None)
        for u in unidades:
            uid = u.get("id_unidade")
            resp_uni       = resp_mapa.get(uid, {})
            u["cpf"]       = mapa.get(uid, {}).get("cpf", "")
            u["email"]     = resp_uni.get("email", "")
            u["endereco"]  = resp_uni.get("endereco", "")
            u["condominio_id"] = id_condominio

    except ValueError as e:
        raise HttpError(404, str(e))
    except Exception as e:
        logger.error(f"Erro ao buscar inadimplentes: {e}")
        raise HttpError(502, "Erro ao consultar Superlógica")

    return 200, {
        "nome_condominio": dados_condo.get("nome_condominio") or f"Condomínio {id_condominio}",
        "dados_condominio": dados_condo,
        "unidades": unidades,
    }


def _parse_body(request):
    import json
    try:
        body = request.body
    except Exception as e:
        logger.error(f"Erro ao ler body da requisição: {e}")
        raise HttpError(400, f"Erro ao ler requisição: {e}")
    try:
        return json.loads(body)
    except Exception as e:
        logger.error(f"Erro ao fazer parse do JSON: {e} | primeiros 200 bytes: {body[:200]}")
        raise HttpError(400, "Corpo da requisição inválido")


def _nome_arquivo(u: dict) -> str:
    nome_safe    = re.sub(r"[^\w\s-]", "", u.get("nome", "desconhecido"))[:40].strip().replace(" ", "_")
    unidade_safe = re.sub(r"[^\w-]", "", u.get("unidade", "uni"))
    return f"{unidade_safe}_{nome_safe}_execucao"


def _extrair_imagem(body: dict) -> bytes | None:
    """Decodifica imagem base64 do campo 'imagem_base64' do body, se presente."""
    import base64
    raw = body.get("imagem_base64") or ""
    if not raw:
        return None
    # Remove prefixo data:image/...;base64,
    if "," in raw:
        raw = raw.split(",", 1)[1]
    try:
        return base64.b64decode(raw)
    except Exception:
        return None


@execucao_router.post("/gerar-docx")
def gerar_docx_endpoint(request):
    """
    Gera DOCX(s) para as unidades selecionadas.
    Uma unidade  → baixa o .docx diretamente.
    Várias unidades → baixa um .zip com todos os .docx.
    """
    body         = _parse_body(request)
    dados_condo  = body.get("condominio", {})
    unidades_sel = body.get("unidades", [])

    if not unidades_sel:
        raise HttpError(400, "Nenhuma unidade selecionada")

    responsavel   = body.get("responsavel") or {}
    modelo        = body.get("modelo") or "com_honorarios"
    m_path        = _modelo_path(modelo)

    if len(unidades_sel) == 1:
        u            = unidades_sel[0]
        imagem_bytes = _extrair_imagem(u)
        variaveis    = _variaveis(dados_condo, u, responsavel, modelo)
        try:
            conteudo = _gerar_docx(variaveis, imagem_bytes=imagem_bytes, modelo_path=m_path)
        except Exception as e:
            logger.error(f"Erro ao gerar DOCX: {e}")
            raise HttpError(500, f"Erro ao gerar documento: {e}")
        nome = f"{_nome_arquivo(u)}.docx"
        resp = HttpResponse(conteudo, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        resp["Content-Disposition"] = f'attachment; filename="{nome}"'
        return resp

    # Múltiplas unidades → ZIP
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for u in unidades_sel:
            imagem_bytes = _extrair_imagem(u)
            variaveis    = _variaveis(dados_condo, u, responsavel, modelo)
            try:
                zf.writestr(f"{_nome_arquivo(u)}.docx", _gerar_docx(variaveis, imagem_bytes=imagem_bytes, modelo_path=m_path))
            except Exception as e:
                logger.error(f"Erro ao gerar DOCX para {u.get('unidade')}: {e}")
    buf.seek(0)
    nome_zip = f"execucao_docx_{date.today().isoformat()}.zip"
    resp = HttpResponse(buf.read(), content_type="application/zip")
    resp["Content-Disposition"] = f'attachment; filename="{nome_zip}"'
    return resp


@execucao_router.post("/gerar-pdf")
def gerar_pdf_endpoint(request):
    """
    Gera PDF(s) para as unidades selecionadas.
    Uma unidade  → baixa o .pdf diretamente.
    Várias unidades → baixa um .zip com todos os .pdf.
    """
    body         = _parse_body(request)
    dados_condo  = body.get("condominio", {})
    unidades_sel = body.get("unidades", [])

    if not unidades_sel:
        raise HttpError(400, "Nenhuma unidade selecionada")

    responsavel   = body.get("responsavel") or {}
    modelo        = body.get("modelo") or "com_honorarios"
    m_path        = _modelo_path(modelo)

    if len(unidades_sel) == 1:
        u            = unidades_sel[0]
        imagem_bytes = _extrair_imagem(u)
        variaveis    = _variaveis(dados_condo, u, responsavel, modelo)
        try:
            conteudo = _gerar_pdf(variaveis, dados_condo.get("nome_condominio", ""), u.get("unidade", ""), imagem_bytes=imagem_bytes, modelo_path=m_path)
        except Exception as e:
            logger.error(f"Erro ao gerar PDF: {e}")
            raise HttpError(500, f"Erro ao gerar documento: {e}")
        nome = f"{_nome_arquivo(u)}.pdf"
        resp = HttpResponse(conteudo, content_type="application/pdf")
        resp["Content-Disposition"] = f'attachment; filename="{nome}"'
        return resp

    # Múltiplas unidades → ZIP
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for u in unidades_sel:
            imagem_bytes = _extrair_imagem(u)
            variaveis    = _variaveis(dados_condo, u, responsavel, modelo)
            try:
                zf.writestr(
                    f"{_nome_arquivo(u)}.pdf",
                    _gerar_pdf(variaveis, dados_condo.get("nome_condominio", ""), u.get("unidade", ""), imagem_bytes=imagem_bytes, modelo_path=m_path)
                )
            except Exception as e:
                logger.error(f"Erro ao gerar PDF para {u.get('unidade')}: {e}")
    buf.seek(0)
    nome_zip = f"execucao_pdf_{date.today().isoformat()}.zip"
    resp = HttpResponse(buf.read(), content_type="application/zip")
    resp["Content-Disposition"] = f'attachment; filename="{nome_zip}"'
    return resp


# ──────────────────────────────────────────────
# CRUD — Responsável pela Petição
# ──────────────────────────────────────────────

@execucao_router.get("/responsaveis-peticao", response={200: list})
def listar_responsaveis(request):
    from core.models import ResponsavelPeticao
    qs = ResponsavelPeticao.objects.all()
    return 200, [
        {"id": r.id, "nome": r.nome, "funcao": r.funcao, "padrao": r.padrao}
        for r in qs
    ]


@execucao_router.post("/responsaveis-peticao", response={201: dict})
def criar_responsavel(request):
    body = _parse_body(request)
    nome   = (body.get("nome") or "").strip()
    funcao = (body.get("funcao") or "").strip()
    padrao = bool(body.get("padrao", False))
    if not nome:
        raise HttpError(400, "Nome é obrigatório")

    from core.models import ResponsavelPeticao
    from django.db import transaction
    with transaction.atomic():
        if padrao:
            ResponsavelPeticao.objects.filter(padrao=True).update(padrao=False)
        r = ResponsavelPeticao.objects.create(nome=nome, funcao=funcao, padrao=padrao)
    return 201, {"id": r.id, "nome": r.nome, "funcao": r.funcao, "padrao": r.padrao}


@execucao_router.put("/responsaveis-peticao/{id_resp}", response={200: dict})
def atualizar_responsavel(request, id_resp: int):
    from core.models import ResponsavelPeticao
    from django.db import transaction
    try:
        r = ResponsavelPeticao.objects.get(pk=id_resp)
    except ResponsavelPeticao.DoesNotExist:
        raise HttpError(404, "Responsável não encontrado")
    body = _parse_body(request)
    nome   = (body.get("nome") or "").strip()
    funcao = (body.get("funcao") or "").strip()
    padrao = bool(body.get("padrao", False))
    if not nome:
        raise HttpError(400, "Nome é obrigatório")
    with transaction.atomic():
        if padrao and not r.padrao:
            ResponsavelPeticao.objects.filter(padrao=True).update(padrao=False)
        r.nome   = nome
        r.funcao = funcao
        r.padrao = padrao
        r.save()
    return 200, {"id": r.id, "nome": r.nome, "funcao": r.funcao, "padrao": r.padrao}


@execucao_router.delete("/responsaveis-peticao/{id_resp}", response={204: None})
def deletar_responsavel(request, id_resp: int):
    from core.models import ResponsavelPeticao
    try:
        r = ResponsavelPeticao.objects.get(pk=id_resp)
    except ResponsavelPeticao.DoesNotExist:
        raise HttpError(404, "Responsável não encontrado")
    r.delete()
    return 204, None


@execucao_router.post("/responsaveis-peticao/{id_resp}/definir-padrao", response={200: dict})
def definir_padrao(request, id_resp: int):
    from core.models import ResponsavelPeticao
    from django.db import transaction
    try:
        r = ResponsavelPeticao.objects.get(pk=id_resp)
    except ResponsavelPeticao.DoesNotExist:
        raise HttpError(404, "Responsável não encontrado")
    with transaction.atomic():
        ResponsavelPeticao.objects.filter(padrao=True).update(padrao=False)
        r.padrao = True
        r.save()
    return 200, {"id": r.id, "nome": r.nome, "funcao": r.funcao, "padrao": r.padrao}
